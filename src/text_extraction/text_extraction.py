import io
import os
import tempfile
from typing import Optional

# For DOCX processing
import docx

# For PDF processing
import fitz  # PyMuPDF
import pytesseract

# For DOC processing (old Word format)
import textract

# For Excel processing
# import pandas as pd
# For HTML processing
from bs4 import BeautifulSoup

# For image processing
from PIL import Image


class DocumentTextExtractor:
    """Extracts text content from various document types.
    Currently supports: PDF, DOCX, DOC, XLSX, XLS, HTML, PNG, JPEG, TXT
    """

    def __init__(self, tesseract_path: Optional[str] = None):
        """Initialize the text extractor.

        Args:
        ----
            tesseract_path: Optional path to Tesseract OCR executable

        """
        # Set Tesseract path if provided
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

        # Map file extensions to extraction methods
        self.extraction_methods = {
            "pdf": self._extract_from_pdf,
            "png": self._extract_from_image,
            "jpg": self._extract_from_image,
            "jpeg": self._extract_from_image,
            "docx": self._extract_from_docx,
            "doc": self._extract_from_doc,
            "html": self._extract_from_html,
            "txt": self._extract_from_txt,
        }
        self.allowed_extensions = set(self.extraction_methods.keys())

    def allowed_file(self, filename: str) -> bool:
        return (
            "." in filename
            and filename.rsplit(".", 1)[1].lower() in self.allowed_extensions
        )

    def extract_text(self, file_data: bytes, filename: str) -> str:
        """Extract text from a file based on its extension

        Args:
        ----
            file_data: Binary file content
            filename: Original filename with extension

        Returns:
        -------
            Extracted text (empty string if extraction failed)

        """
        # Get file extension
        file_extension = os.path.splitext(filename)[1].lower().lstrip(".")

        # Check if we support this file type
        if file_extension not in self.extraction_methods:
            return f"Unsupported file type: {file_extension}"

        # Call the appropriate extraction method
        return self.extraction_methods[file_extension](file_data)

    def _extract_from_pdf(self, file_data: bytes) -> str:
        """Extract text from PDF files using PyMuPDF"""
        try:
            text = ""

            with fitz.open(stream=file_data, filetype="pdf") as doc:
                # Extract text from each page
                for page in doc:
                    text += page.get_text() + "\n"

            return text

        except Exception as e:
            return f"PDF extraction error: {str(e)}"

    def _extract_from_image(self, file_data: bytes) -> str:
        """Extract text from images using OCR"""
        try:
            # Open image using PIL
            image = Image.open(io.BytesIO(file_data))

            # Perform OCR to extract text
            text = pytesseract.image_to_string(image)

            return text

        except Exception as e:
            return f"Image extraction error: {str(e)}"

    def _extract_from_docx(self, file_data: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            # Create a temporary file-like object
            docx_file = io.BytesIO(file_data)

            # Open with python-docx
            doc = docx.Document(docx_file)

            # Extract all paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Extract tables content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"

            return text

        except Exception as e:
            return f"DOCX extraction error: {str(e)}"

    def _extract_from_doc(self, file_data: bytes) -> str:
        """Extract text from DOC files using textract"""
        try:
            # Create a temporary file to work with
            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                temp_filename = temp_file.name
                temp_file.write(file_data)

            # Extract text using textract
            text = textract.process(temp_filename).decode("utf-8", errors="replace")

            # Clean up
            os.unlink(temp_filename)

            return text

        except Exception as e:
            return f"DOC extraction error: {str(e)}"

    def _extract_from_html(self, file_data: bytes) -> str:
        """Extract text from HTML files"""
        try:
            # Decode the HTML content
            html_content = file_data.decode("utf-8", errors="replace")

            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_content, "html.parser")

            # Extract text
            text = soup.get_text(separator="\n", strip=True)

            return text

        except Exception as e:
            return f"HTML extraction error: {str(e)}"

    def _extract_from_txt(self, file_data: bytes) -> str:
        """Extract text from plain text files"""
        try:
            # Decode the text content
            text = file_data.decode("utf-8", errors="replace")

            return text

        except Exception as e:
            return f"Text file extraction error: {str(e)}"
